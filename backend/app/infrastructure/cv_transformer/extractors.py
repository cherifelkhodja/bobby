"""Text extractors for PDF and DOCX files.

Implements CvTextExtractorPort for dependency inversion.
"""

import io
import logging
from typing import BinaryIO

from docx import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# DOCX XML namespaces for text box extraction
_WPS_TXBX = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
_VML_TEXTBOX = "{urn:schemas-microsoft-com:vml}textbox"
_W_P = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
_W_T = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"


class PdfTextExtractor:
    """PDF text extractor implementing CvTextExtractorPort."""

    def extract(self, content: bytes) -> str:
        """Extract text from PDF content."""
        return extract_text_from_pdf(content)


class DocxTextExtractor:
    """DOCX text extractor implementing CvTextExtractorPort."""

    def extract(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        return extract_text_from_docx(content)


def extract_text_from_pdf(file_content: bytes | BinaryIO) -> str:
    """Extract text content from a PDF file.

    Args:
        file_content: PDF file content as bytes or file-like object.

    Returns:
        Extracted text from all pages.

    Raises:
        ValueError: If the PDF cannot be read or is empty.
    """
    if isinstance(file_content, bytes):
        file_content = io.BytesIO(file_content)

    try:
        reader = PdfReader(file_content)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        if not text_parts:
            raise ValueError("Le PDF ne contient pas de texte extractible")

        return "\n\n".join(text_parts)
    except Exception as e:
        if "texte extractible" in str(e):
            raise
        raise ValueError(f"Erreur lors de la lecture du PDF: {str(e)}")


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Extract text from file content based on file extension.

    Args:
        content: File content as bytes.
        filename: Original filename to determine file type.

    Returns:
        Extracted text.

    Raises:
        ValueError: If extraction fails or file type is not supported.
    """
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(content)
    else:
        raise ValueError(f"Format de fichier non supporté: {filename}")


def _extract_textbox_text(document: Document) -> list[str]:
    """Extract text from text boxes and shapes in a DOCX document.

    Many CV templates store content in floating text boxes which are not
    included in document.paragraphs or document.tables.

    Handles:
    - wps:txbx  (Word Processing Shapes — modern DOCX text boxes)
    - v:textbox (VML — older format text boxes)
    """
    text_parts: list[str] = []

    for container_tag in (_WPS_TXBX, _VML_TEXTBOX):
        for container in document.element.iter(container_tag):
            para_texts: list[str] = []
            for para in container.iter(_W_P):
                t_nodes = [t.text for t in para.iter(_W_T) if t.text]
                para_text = "".join(t_nodes).strip()
                if para_text:
                    para_texts.append(para_text)
            if para_texts:
                text_parts.append("\n".join(para_texts))

    return text_parts


def extract_text_from_docx(file_content: bytes | BinaryIO) -> str:
    """Extract text content from a DOCX file.

    Args:
        file_content: DOCX file content as bytes or file-like object.

    Returns:
        Extracted text from all paragraphs, tables and text boxes.

    Raises:
        ValueError: If the DOCX cannot be read or is empty.
    """
    if isinstance(file_content, bytes):
        file_content = io.BytesIO(file_content)

    try:
        document = Document(file_content)
        text_parts = []

        # Extract text from paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        # Extract text from text boxes / shapes (wps:txbx, v:textbox).
        # Many CV templates use floating text boxes that are invisible to
        # document.paragraphs and document.tables.
        textbox_parts = _extract_textbox_text(document)
        if textbox_parts:
            logger.debug(
                "DOCX: extracted %d text box(es) in addition to %d paragraph/table block(s)",
                len(textbox_parts),
                len(text_parts),
            )
            text_parts.extend(textbox_parts)

        if not text_parts:
            raise ValueError("Le document Word ne contient pas de texte extractible")

        return "\n\n".join(text_parts)
    except Exception as e:
        if "texte extractible" in str(e):
            raise
        raise ValueError(f"Erreur lors de la lecture du document Word: {str(e)}")
